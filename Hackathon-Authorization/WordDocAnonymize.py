import os
import re
import spacy

from urllib.parse import urlparse

from extract_msg import Message

from pathlib import Path

import email
import pdfkit
from bs4 import BeautifulSoup  
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from docx import Document

import random

user_map = {}
domain_map = {}
email_map = {}

people = set()
# orgs = set()

exclude_people = (
    'lim', 'li', 'ratio', 'prem', 'reinst', 'premium', 'underwriter', 'order', 'draft', 'rplu',
    'suite', 'lang', 'gucng', 'gte', 'generator', 'mso', 'word', 'xl', 'casualty', 'slips', 'name',
    'ark', 'syndicate', 'property', 'synd', 'trident', 'miami', 'green', 'shield', 'greenshield', 'canada',
    'summary', 'street', 'ste', 'tel', 'nursing', 'loss', 'jlt', 'box', 'exhibits', 'askew', 'gbr', 'fwd', 'sent', 'model',
    'request', 'webdings', 'response', 'canadas', 'duty', 'nowrap', 'terms', 'div', 'terms', 'tab', 'plan', 'gsc', 'everest',
    'everestre', 'designs', 'meta', 'liabilities', 'canadian', 'rue', 'form', 'notification', 'planholder', 'plaza',
    'details', 'assumed', 'document', 'broker', 'primary', 'cyber', 'cancel', 'conract', 'assumed', 'primary', 'remarks',
    'term', 'type', 'mep', 'basis', 'reinsured', 'basis', 'catastrope', 'occurrence', 'calamity', 'damage', 'insurer',
    'jurisdiction', 'nett', 'reinsurance', 'reinsurer', 'trojan', 'wording', 'contract', 'top', 'layer', 'addedby', 'max',
    'worldwide', 'catastrophe', 'reference', 'status', 'limit', 'corporation', 'title', 'endorsement', 'elr', 'agg', 'ezflow', 'sum',
    'llowcvcr', 'usds1', 'tios', 'continuous', 'party', 'agreement', 'building', 'contents', 'regions', 'deductible', 'continuous', 'convective',
    'worldwide', 'following', 'contact', 'models', 'misdescription', 'llowcvcr', 'contact', 'fhcf', 'outgo', 'square', 'footage', 'transfer', 'catastrophe',
    'fungi', 'retrieval', 'modeling', 'nil', 'iso', 'slip', 'usds1', 'USDSl', 'geometry', 'structures', 'leakage', 'wind', 'leakage', 'format', 'misc', 'surge',
    'retention', 'tsunami', 'portfolio', 'data', 'complex', 'fire', 'acceptance', 'cecr', 'utilized', 'demand', 'waiver', 'all', 'winter', 'smoke', 'covering', 'gal',
    'exclusion', 'construnction', 'renew', 'ocean', 'premiums', 'upper', 'surplus', 'share', 'main', 'offshore', 'occ', 'sudden', 'roof'
    )

counters = {
    "domain_counter": 0,
}

names = [
    "Emma", "Liam", "Olivia", "Noah", "Ava", "Ethan", "Sophia", "Mason", 
    "Isabella", "Lucas", "Mia", "Henry", "Charlotte", "Alexander", "Amelia", 
    "Michael", "Harper", "Benjamin", "Evelyn", "Daniel", "Aria", "Matthew", 
    "Scarlett", "James", "Grace", "William", "Chloe", "Jack", "Victoria", 
    "Luke", "Riley", "David", "Lily", "Joseph", "Ellie", "Thomas", "Hannah", 
    "Samuel", "Julia", "Andrew", "Layla", "Gabriel", "Natalie", "John", 
    "Zoey", "Christopher", "Penelope", "Joshua", "Mila", "Ryan",  
    "Miles", "Sophie", "Derek", "Tanya", "Omar", "Elaine", "Caleb", "Isla", "Rafael", "Bianca",
    "Liam", "Priya", "Jonas", "Wendy", "Nina", "Mateo", "Jared", "Carmen", "Leo", "Gianna",
    "Xander", "Maya", "Andres", "Natalie", "Hugo", "Tessa", "Silas", "Camila", "Felix", "Zara",
    "Lucia", "Arjun", "Clara", "Malik", "Daisy", "Roman", "Amira", "Elias", "Avery", "Victor",
    "Noor", "Ezekiel", "Luna", "Griffin", "Heidi", "Enzo", "Jasmine", "Theo", "Sadie", "Nico",
    "Ruby", "Axel", "Freya", "Owen", "Elena", "Soren", "Chloe", "Micah", "Lara", "Jude",
    "Anya", "Cyrus", "Phoebe", "Ezra", "Ivy", "Marco", "Layla", "Zane", "Sienna", "Remy",
    "Amina", "Gavin", "Skye", "Tristan", "Nova", "Lorenzo", "Esme", "Cassius", "Vivian", "Jax",
    "Mira", "Damien", "Elodie", "Harvey", "Keira", "Quinn", "Rocco", "Alina", "Emilio", "Talia",
    "Brody", "Serena", "Kai", "Maeve", "Dallas", "Yara", "Colt", "Juniper", "Reed", "Nadia",
    "Aaliyah", "Aaron", "Abigail", "Adam", "Addison", "Adrian", "Adriana", "Aidan", "Ainsley", "Alan",
    "Alana", "Albert", "Alex", "Alexa", "Alexander", "Alexandra", "Alexis", "Alfred", "Alice", "Alicia",
    "Alisha", "Allison", "Alma", "Alonzo", "Alton", "Alvin", "Amanda", "Amber", "Amelia", "Amos",
    "Amy", "Ana", "Andre", "Andrea", "Andrew", "Angel", "Angela", "Angelica", "Angelo", "Angus",
    "Anita", "Ann", "Anna", "Annabelle", "Anthony", "Antoinette", "Antonio", "April", "Archer", "Ariana",
    "Arianna", "Ariel", "Arielle", "Arlo", "Armando", "Arnold", "Arthur", "Arturo", "Ashley",
    "Ashton", "Astrid", "Aubrey", "Audrey", "August", "Augustus", "Aurora", "Austin", "Autumn", "Ava",
    "Avery", "Bailey", "Barbara", "Barry", "Beatrice", "Beau", "Bella", "Ben", "Benjamin", "Bennett",
    "Bernadette", "Bernard", "Bertha", "Beth", "Bethany", "Betty", "Bianca", "Bill", "Billy", "Blair",
    "Blake", "Blanche", "Bob", "Bobby", "Bonnie", "Brad", "Bradley", "Brady", "Brandon", "Brenda",
    "Brendan", "Brent", "Brett", "Brian", "Brianna", "Bridget", "Brittany", "Brock", "Brody", "Bruce",
    "Bryan", "Bryce", "Caleb", "Callie", "Cameron", "Camila", "Camille", "Carl", "Carla", "Carlos",
    "Carly", "Carmen", "Carol", "Caroline", "Carolyn", "Carrie", "Carson", "Carter", "Casey", "Cassandra",
    "Cassidy", "Catherine", "Cecilia", "Cecil", "Celeste", "Celia", "Chad", "Chance", "Chandler", "Charles",
    "Charlotte", "Chase", "Chelsea", "Cheryl", "Chester", "Chloe", "Chris", "Christian", "Christina", "Christine",
    "Christopher", "Chuck", "Ciara", "Cindy", "Claire", "Clara", "Clarence", "Clark", "Claudia", "Clayton",
    "Clement", "Clementine", "Clifford", "Clifton", "Clint", "Cody", "Colby", "Cole", "Colin", "Collin",
    "Connie", "Connor", "Constance", "Cooper", "Cora", "Corey", "Corinne", "Cortez", "Courtney", "Craig",
    "Crystal", "Curtis", "Cynthia", "Daisy", "Dale", "Dallas", "Dalton", "Damian", "Damien", "Dana",
    "Daniel", "Daniela", "Danielle", "Danny", "Daphne", "Darcy", "Daria", "Darin", "Darius", "Darla",
    "Darrell", "Darren", "Dave", "David", "Dawn", "Dean", "Deanna", "Debbie", "Deborah", "Declan",
    "Dee", "Deirdre", "Delbert", "Delia", "Dell", "Delores", "Delroy", "Demi", "Dennis", "Denny",
    "Derek", "Derrick", "Desiree", "Desmond", "Devin", "Devon", "Diana", "Diane", "Diego", "Dillon",
    "Dirk", "Dixie", "Dominic", "Dominique", "Don", "Donald", "Donna", "Dora", "Doris", "Dorothy",
    "Doug", "Douglas", "Drew", "Duane", "Dustin", "Dwight", "Dylan", "Earl", "Easton", "Ebony",
    "Ed", "Eddie", "Eden", "Edgar", "Edith", "Edmund", "Edna", "Eduardo", "Edward", "Edwin",
    "Effie", "Eileen", "Elaine", "Eli", "Elias", "Elijah", "Elise", "Eliza", "Elizabeth", "Ella",
    "Ellen", "Ellie", "Elliot", "Elliott", "Ellis", "Elmer", "Elsie", "Elton", "Elvis", "Emily",
    "Emma", "Emmanuel", "Emmett", "Eric", "Erica", "Erick", "Erin", "Ernest", "Ernie", "Esmeralda",
    "Esteban", "Esther", "Ethan", "Ethel", "Eugene", "Eva", "Evan", "Eve", "Evelyn", "Everett",
    "Ezekiel", "Ezra", "Fabian", "Faith", "Fallon", "Fannie", "Farrah", "Fatima", "Felicia", "Felix",
    "Fernando", "Finn", "Fiona", "Floyd", "Frances", "Francesca", "Francis", "Frank", "Frankie", "Franklin",
    "Fred", "Freddie", "Frederick", "Freya", "Gabriel", "Gabriela", "Gabriella", "Gage", "Gail", "Gale",
    "Garrett", "Garrison", "Gary", "Gavin", "Gene", "Genevieve", "Geoffrey", "George", "Georgia", "Gerald",
    "Geraldine", "Gerard", "Gertrude", "Gia", "Gianna", "Gideon", "Gilbert", "Gilda", "Gillian", "Gina",
    "Ginger", "Giovanni", "Glen", "Glenda", "Gloria", "Gordon", "Grace", "Gracie", "Grant", "Greg",
    "Gregg", "Gregory", "Greta", "Griffin", "Griselda", "Guadalupe", "Gwendolyn", "Gustavo", "Guy", "Hadley",
    "Hailey", "Haley", "Hallie", "Hank", "Hannah", "Hans", "Harley", "Harold", "Harper", "Harris",
    "Harrison", "Harry", "Harvey", "Hazel", "Heather", "Heidi", "Helen", "Helena", "Henry", "Herbert",
    "Herman", "Hester", "Hilda", "Hillary", "Holly", "Homer", "Hope", "Horace", "Howard", "Hubert",
    "Hugh", "Hugo", "Hunter", "Ian", "Ida", "Ignacio", "Imani", "Imelda", "India", "Ingrid",
    "Ira", "Irvin", "Irving", "Isaac", "Isabel", "Isabella", "Isaiah", "Isidro", "Isla", "Israel",
    "Ivy", "Jack", "Jackie", "Jackson", "Jaclyn", "Jacob", "Jacqueline", "Jade", "Jaden", "Jaime",
    "Jalen", "Jamal", "James", "Jamie", "Jan", "Jana", "Jane", "Janelle", "Janet", "Janice",
    "Jared", "Jasmine", "Jason", "Jasper", "Javier", "Jay", "Jayden", "Jayla", "Jaylen", "Jeff",
    "Jefferson", "Jeffrey", "Jenna", "Jennifer", "Jenny", "Jenson", "Jeremy", "Jerome", "Jerry", "Jesse",
    "Jessica", "Jessie", "Jesus", "Jill", "Jillian", "Jim", "Jimmy", "Joan", "Joanna", "Jocelyn",
    "Jodi", "Jodie", "Jody", "Joe", "Joel", "Joey", "John", "Johnathan", "Johnnie", "Johnny",
    "Jolene", "Jon", "Jonah", "Jonathan", "Joni", "Jordan", "Jordyn", "Jorge", "Jose", "Joseph",
    "Josephine", "Josh", "Joshua", "Josiah", "Josselyn", "Joy", "Joyce", "Juan", "Juana", "Judah",
    "Judd", "Jude", "Judith", "Judy", "Julia", "Julian", "Juliana", "Julianna", "Julie", "Juliet",
    "Julio", "Julius", "June", "Juniper", "Justina", "Justin", "Justine", "Kaden", "Kai", "Kaia",
    "Kaila", "Kaitlin", "Kaleb", "Kali", "Kallie", "Kamari", "Kamila", "Kanye", "Kara", "Karen",
    "Kari", "Karin", "Karina", "Karl", "Karla", "Karlee", "Karly", "Kasey", "Kash", "Kasia",
    "Kate", "Katelyn", "Katelynn", "Katerina", "Katharine", "Katherine", "Kathleen", "Kathryn", "Kathy", "Katie",
    "Katrina", "Kay", "Kaya", "Kayden", "Kayla", "Kaylee", "Kayleigh", "Kazim", "Keanu", "Keaton",
    "Keegan", "Keisha", "Keith", "Kellan", "Kellen", "Kelly", "Kelsey", "Kelvin", "Ken", "Kendall",
    "Kendra", "Kendrick", "Kenji", "Kenna", "Kenneth", "Kenny", "Kent", "Kenya", "Kermit", "Kerry",
    "Kevin", "Khalid", "Kian", "Kiara", "Kiera", "Kieran", "Kim", "Kimberly", "King", "Kingsley",
    "Kirk", "Klaus", "Knox", "Kobe", "Kody", "Kolby", "Konstantin", "Kora", "Kris", "Kristen",
    "Kristi", "Kristian", "Kristie", "Kristin", "Kristina", "Kristopher", "Kurt", "Kurtis", "Kyla",
    "Kyle", "Kylee", "Kyler", "Kyra", "Lacey", "Lacy", "Ladonna", "Laisha", "Lakisha", "Lamar",
    "Lamont", "Lance", "Lane", "Lara", "Larry", "Latasha", "Laura", "Lauren", "Laurence", "Laurie",
    "Lawanda", "Lawrence", "Layla", "Leah", "Leandro", "Lee", "Leif", "Leila", "Leilani", "Leia",
    "Lena", "Lenard", "Lenny", "Leo", "Leona", "Leonard", "Leroy", "Lesley", "Leslie", "Lester",
    "Levi", "Lewis", "Liam", "Lila", "Lilian", "Liliana", "Lillian", "Lily", "Lin", "Linda",
    "Lindsay", "Lindsey", "Lisa", "Liz", "Lizbeth", "Lizzie", "Lloyd", "Logan", "Lois", "Lola",
    "Lorena", "Lorenzo", "Lori", "Lorna", "Louis", "Louisa", "Louise", "Lucas", "Lucia", "Lucian",
    "Luciana", "Lucille", "Lucinda", "Lucky", "Lucy", "Luis", "Luisa", "Luke", "Luna", "Luther",
    "Lydia", "Lyle", "Lynda", "Lynn", "Lynne", "Mabel", "Macey", "Macie", "Mackenzie", "Macy",
    "Madalyn", "Maddie", "Madison", "Madeline", "Madelyn", "Mae", "Magdalena", "Maggie", "Magnus", "Mahalia",
    "Maia", "Major", "Makayla", "Makenna", "Makenzie", "Malachi", "Malcolm", "Mallory", "Malvina", "Mandy",
    "Manuel", "Mara", "Marc", "Marcel", "Marcela", "Marcelo", "Marcia", "Marco", "Marcos", "Marcus",
    "Margaret", "Margarita", "Margie", "Margo", "Maria", "Mariah", "Mariam", "Marian", "Mariana", "Marianna",
    "Marie", "Marilyn", "Marina", "Mario", "Marion", "Marisa", "Marisol", "Mark", "Markus", "Marla",
    "Marlon", "Marsha", "Marshall", "Martha", "Martin", "Martina", "Marty", "Marvin", "Mary", "Mason",
    "Mateo", "Mathew", "Mathias", "Matt", "Matthew", "Matthias", "Maud", "Maureen", "Maurice", "Max",
    "Maximilian", "Maxwell", "May", "Maya", "Mckayla", "Mckenna", "Mckenzie", "Meagan", "Megan", "Meghan",
    "Melanie", "Melinda", "Melissa", "Melvin", "Mercedes", "Meredith", "Merle", "Mia", "Micah", "Michael",
    "Michaela", "Michelle", "Miguel", "Mike", "Mila", "Miles", "Millard", "Millie", "Milo", "Milton",
    "Mina", "Minnie", "Mira", "Miriam", "Misty", "Mitchell", "Molly", "Mona", "Monica", "Monique",
    "Monty", "Morgan", "Morris", "Moses", "Muriel", "Murray", "Myah", "Mya", "Myles", "Myra",
    "Naomi", "Nash", "Natalia", "Natalie", "Natasha", "Nathan", "Nathanael", "Neal", "Ned", "Neil",
    "Nell", "Nellie", "Nelson", "Neva", "Nevaeh", "Neville", "Newton", "Nicholas", "Nick", "Nicole",
    "Nigel", "Nikki", "Nina", "Noah", "Noel", "Noelle", "Nola", "Nolan", "Norma", "Norman",
    "Nova", "Octavia", "Octavio", "Olaf", "Olga", "Oliver", "Olivia", "Ollie", "Omar", "Ora",
    "Orlando", "Oscar", "Osvaldo", "Otis", "Otto", "Owen", "Paige", "Pamela", "Paolo", "Paris",
    "Parker", "Patricia", "Patrick", "Patty", "Paul", "Paula", "Pauline", "Paxton", "Payton", "Pearl",
    "Pedro", "Penny", "Percy", "Perry", "Pete", "Peter", "Petra", "Phil", "Philip", "Phillip",
    "Phoebe", "Phyllis", "Pierce", "Pierre", "Piper", "Pippa", "Porter", "Preston", "Prince", "Priscilla",
    "Quentin", "Quincy", "Quinn", "Quinton", "Rachel", "Rafael", "Raheem", "Raina", "Ralph", "Ramiro",
    "Ramona", "Ramon", "Randall", "Randy", "Raphael", "Raquel", "Rashad", "Raul", "Raven", "Ray",
    "Raymond", "Rayna", "Reagan", "Rebecca", "Rebekah", "Reed", "Reese", "Regan", "Reggie", "Regina",
    "Reginald", "Reid", "Reilly", "Remy", "Rene", "Renee", "Rex", "Rey", "Reyes", "Reynaldo",
    "Rhiannon", "Ricardo", "Rich", "Richard", "Rick", "Ricky", "Rihanna", "Riley", "Rita", "River",
    "Robert", "Roberta", "Roberto", "Robin", "Rocco", "Rochelle", "Rocky", "Rod", "Rodney", "Rodolfo",
    "Roger", "Roland", "Rolando", "Rolf", "Roma", "Roman", "Romeo", "Ron", "Ronald", "Ronnie",
    "Rosa", "Rosalie", "Rosalind", "Rosalyn", "Rose", "Rosemary", "Ross", "Rowan", "Roxanne", "Roy",
    "Royce", "Ruby", "Rudolph", "Rudy", "Russell", "Ruth", "Ryan", "Ryann", "Ryder", "Rylan",
    "Sabrina", "Sadie", "Sage", "Sal", "Salena", "Sally", "Salvador", "Sam", "Samantha", "Samir",
    "Samuel", "Sandra", "Sandy", "Santiago", "Sara", "Sarah", "Sasha", "Saul", "Savannah", "Sawyer",
    "Scarlett", "Scott", "Sean", "Sebastian", "Selena", "Selma", "Sergio", "Seth", "Shaina", "Shane",
    "Shaniqua", "Shannon", "Shari", "Sharon", "Shaun", "Shawn", "Shawna", "Shea", "Sheila", "Shelby",
    "Sheldon", "Shelley", "Sheri", "Sherman", "Sherri", "Sherry", "Shirley", "Sidney", "Siena", "Sierra",
    "Silas", "Silvia", "Simon", "Simone", "Sky", "Skylar", "Skyler", "Sloane", "Sofia", " സോഫിയ", "Sol",
    "Solomon", "Sonia", "Sonny", "Sophia", "Sophie", "Spencer", "Stacey", "Stacy", "Stan", "Stanley",
    "Starla", "Stefan", "Stefanie", "Stephen", "Stephanie", "Sterling", "Steve", "Steven", "Stevie", "Stewart",
    "Stone", "Stuart", "Summer", "Sunny", "Susan", "Susana", "Susie", "Suzanne", "Sydney", "Williamson",
    "Jasper", "Lena", "Miles", "Zara", "Orion", "Ivy", "Dante", "Talia", "Felix", "Maya",
    "Silas", "Nina", "Ezekiel", "Rhea", "Leo", "Sienna", "Arlo", "Quinn", "Ezra", "Layla",
    "Kai", "Nova", "Beckett", "Freya", "Axel", "Celeste", "Otis", "Vera", "Ronan", "Thea",
    "Mateo", "Skye", "Cassian", "Juniper", "Xander", "Mabel", "Finn", "Dahlia", "Orla", "Theo",
    "Rowan", "Esme", "Cyrus", "Luna", "Emmett", "Iris", "Zane", "Cleo", "Jude", "Sage",
    "Caleb", "Aria", "Rory", "Hazel", "Damien", "Isla", "Soren", "Elsie", "Luca", "Wren",
    "Atlas", "Phoebe", "Gideon", "Kaia", "Nico", "Tess", "Hugo", "Ayla", "Tobias", "Maeve",
    "Asher", "Noa", "Bodhi", "Lilah", "Levi", "Indie", "Graham", "Faye", "Declan", "Lumi",
    "Kieran", "Anya", "Eli", "Zadie", "Reed", "Nola", "Simon", "Gaia", "Milo", "Calla",
    "Ambrose", "Willa", "Jaxon", "Demi", "Rafael", "Petra", "Blaise", "Ada", "Griffin", "Clio",
    "Thalia", "Cormac", "Saskia", "Anders", "Maren", "Boone", "Eleni", "Jago", "Aviva", "Renzo",
    "Briony", "Callahan", "Isolde", "Stellan", "Zinnia", "Torin", "Noemi", "Bastian", "Odette", "Keir",
    "Sabine", "Rafe", "Lilith", "Viggo", "Bronte", "Ewan", "Amaris", "Dorian", "Yara", "Halston",
    "Tamsin", "Lucien", "Maris", "Baylor", "Seren", "Jorah", "Imani", "Cairo", "Delphine", "Evander",
    "Neriah", "Leif", "Paloma", "Roux", "Colm", "Calista", "Merrick", "Sable", "Auden", "Zuleika",
    "Lazaro", "Ines", "Koa", "Phaedra", "Bran", "Galatea", "Thorne", "Elara", "Magnus", "Azura",
    "Fintan", "Oona", "Cassius", "Aisling", "Quillon", "Xiomara", "Greer", "Idris", "Vada", "Ansel",
    "Tova", "Dax", "Alouette", "Hale", "Sunniva", "Zen", "Milou", "Ivar", "Thisbe", "Oren",
    "Elowen", "Kenzo", "Liora", "Phineas", "Bryn", "Oswin", "Alethea", "Roan", "Yseult", "Zephyr",
    "Alaric", "Maelis", "Eirik", "Sunnie", "Jovan", "Twila", "Nevin", "Lazuli", "Oisin", "Mireille",
    "Torsten", "Alba", "Rian", "Petal", "Casimir", "Navi", "Jem", "Sabriel", "Orin", "Zaria",
    "Caelum", "Drea", "Fenris", "Lilou", "Kian", "Sybil", "Thorne", "Ismay", "Ephraim", "Zosia",
    "Lorcan", "Tirzah", "Matisse", "Anouk", "Quade", "Idalia", "Bram", "Sunniva", "Leonie", "Ashby",
    "Verne", "Noor", "Caspian", "Lyra", "Bex", "Elian", "Sorrel", "Talon", "Amira", "Nox",
    "Cedric", "Odile", "Harlan", "Vesper", "Arden", "Blythe", "Perrin", "Zeva", "Rowe", "Emrys",
    "Cyan", "Delia", "Jules", "Ariadne", "Oberon", "Mira", "Koa", "Yuna", "Eira", "Halcyon",
    "Ione", "Azriel", "Greta", "Stellan", "Alaric", "Lilia", "Boaz", "Nyra", "Frey", "Romilly",
    "Galen", "miles.milo"
]

len_names = len(names)

email_addresses = [
    "emma.watson", "liam.smith", "olivia.jones", "noah.brown", "ava.taylor", 
    "ethan.wilson", "sophia.moore", "mason.jackson", "isabella.martin", "lucas.thomas", 
    "mia.white", "henry.harris", "charlotte.lewis", "alexander.clark", "amelia.walker", 
    "michael.hall", "harper.allen", "benjamin.young", "evelyn.king", "daniel.wright", 
    "aria.scott", "matthew.green", "scarlett.adams", "james.baker", "grace.nelson", 
    "william.carter", "chloe.mitchell", "jack.roberts", "victoria.turner", "luke.phillips", 
    "riley.campbell", "david.parker", "lily.evans", "joseph.edwards", "ellie.collins", 
    "thomas.stewart", "hannah.morris", "samuel.rogers", "julia.reed", "andrew.cook", 
    "layla.morgan", "gabriel.murphy", "natalie.bailey", "john.cooper", "zoey.ward", 
    "christopher.bell", "penelope.rivera", "joshua.cox", "mila.brooks", "ryan.kelly", 
    "sophie.bennett", "caleb.howard", "madison.perry", "isaac.long", "hazel.foster", 
    "nathan.griffin", "aubrey.sanders", "logan.price", "claire.russell", "eli.henderson", 
    "violet.graham", "dylan.ross", "luna.dixon", "jackson.watson", "stella.hunter", 
    "owen.mcdonald", "emily.hicks", "levi.palmer", "aurora.barnes", "gavin.sullivan", 
    "addison.wells", "carson.fisher", "piper.hamilton", "miles.hudson", "ruby.fleming", 
    "asher.stone", "brooklyn.webb", "declan.ellis", "eliana.richards", "ryder.dunn", 
    "savannah.hart", "parker.holmes", "audrey.mason", "tristan.warren", "delilah.hawkins", 
    "ezra.dean", "freya.matthews", "jonah.lawson", "ivy.simmons", "micah.walters", 
    "esther.crawford", "silas.burton", "lydia.ford", "abraham.daniels", "marie.garrett", 
    "finley.horton", "june.spencer", "colton.miles", "eden.franklin", "miles.milo"
]

len_email_addresses = len(email_addresses)

replacements = {
    "BMG": "Whiteboard",
    "SEGUROS": "Enterprises",
    "Austral": "Brick",
    "Seguaradora": "House",
    "MITSUI": "White",
    "SUMITOMO": "Oak Inc",
    "Matheus": "Mark",
    "Sales": "Blake",
    "Roberto": "Megan",
    "Flavia": "Fluer", 
    "Hammerle": "Rajan",
    "Ark": "Sunset",
    "Syndicate": "Sky",
    "Synd": "Sky",
    "Trident": "Ocean Storm",
    "TRIDENTINSURANCECO": "OceanStormInsuranceCo",
    'Green': 'Solid',
    'Shield': 'Earth',
    'GreenShield': 'SolidEarth'
}

def decode_if_bytes(data, encoding='utf-8'):
    if isinstance(data, bytes):
        return data.decode(encoding, errors='replace')
    return data

def obfuscate_email_addresses_structured(text):
    if text is None:
        return text
    
    # Email pattern
    email_pattern = r'([a-zA-Z0-9_.+-]+)@([a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)'
    
    # Find all unique emails
    matches = re.findall(email_pattern, text)
    unique_emails = sorted(set(f"{local}@{domain}" for local, domain in matches))
    
    for email in unique_emails:
        local, domain = email.split("@")

        if local not in user_map:
            rand_email_index = random.randint(0, len_email_addresses - 1)
            # print(rand_email_index)
            user_map[local] = email_addresses[rand_email_index]

        if domain not in domain_map:
            domain_map[domain] = f"domain{counters["domain_counter"]}.com"
            counters["domain_counter"] = counters["domain_counter"] + 1

        obfuscated_email = f"{user_map[local]}@{domain_map[domain]}"
        email_map[email] = obfuscated_email

    # Replace emails in text
    def replace_email(match):
        full_email = f"{match.group(1)}@{match.group(2)}"
        return email_map.get(full_email, full_email)

    obfuscated_text = re.sub(email_pattern, replace_email, text)
    
    return obfuscated_text

def obfuscate_entities(text):
    if text is None:
        return text
    
    nlp = spacy.load("en_core_web_sm")

    text = obfuscate_email_addresses_structured(text)

    doc = nlp(text)

    # Extract named entities
    for ent in doc.ents:
        # print(f'ent.label_: {ent.label_}, ent.text:{ent.text}')
        if ent.label_ == "PERSON":
            name_words = ent.text.split()
            if len(name_words) > 5:
                continue
            for name_word in name_words:
                if name_word.isalpha() and name_word.lower() not in exclude_people and len(name_word) > 2:
                    people.add(name_word)
        # elif ent.label_ == "ORG":
        #     orgs.add(ent.text)

        # print('people:', people)

    # Create mappings for obfuscation
    obfuscation_map = {}
    
    for idx, name in enumerate(sorted(people), 1):
        rand_name_index = random.randint(0, len_names - 1)
        # print(f'Index: {idx}, name:{name}, rand_name_index:{rand_name_index}')
        obfuscation_map[name] = names[rand_name_index]
    # for idx, org in enumerate(sorted(orgs), 1):
    #     obfuscation_map[org] = f"Org{idx}"


    # Combine all patterns into one regex
    combined_pattern = re.compile(
        "|".join(re.escape(k) for k in sorted(obfuscation_map, key=lambda x: -len(x))), 
        flags=re.IGNORECASE
    )

    # Replace function
    def replace_match(match):
        matched = match.group(0)
        for original, obfuscated in obfuscation_map.items():
            if matched.lower() == original.lower():
                return obfuscated
        return matched

    # Replace all entities
    obfuscated_text = combined_pattern.sub(replace_match, text)

    for old_name, new_name in replacements.items():
        obfuscated_text = re.sub(re.escape(old_name), new_name, obfuscated_text, flags=re.IGNORECASE)

    return obfuscated_text

def anonymize_and_write_docs(input_path, output_path):
    doc = Document(input_path)

    for para in doc.paragraphs:
        text = decode_if_bytes(para.text)
        # print(text)
        para.text = obfuscate_entities(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = obfuscate_entities(cell.text)

    # Headers and footers
    for idx, section in enumerate(doc.sections, 0):
        print(idx)
        header = section.header
        footer = section.footer

        for para in header.paragraphs:
            para.text = obfuscate_entities(para.text)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = obfuscate_entities(cell.text)

        for para in footer.paragraphs:
            para.text = obfuscate_entities(para.text)

        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = obfuscate_entities(cell.text)

    # Save the new document file
    output_path = obfuscate_entities(output_path)
    directory = os.path.dirname(output_path)
    print(f'MakeDir: {directory}')
    os.makedirs(directory, exist_ok=True)
    print(output_path)

    doc.save(output_path)
    print(f"Modified document saved to {output_path}")


def get_files_with_extension(directory, extension):
    return [str(file) for file in Path(directory).rglob(f'*{extension}')]

def anonymize_docs():
    input_directory = "C:\\work\\data\\original"
    output_directory = "C:\\work\\data\\anonymized"
    extension = ".msg"

    input_files = get_files_with_extension(input_directory, extension)
    print('people:', people)
    for input_file in input_files:
        output_eml_file = input_file.replace(input_directory, output_directory)
        output_eml_file = output_eml_file.replace(extension, '.eml')
        print(input_file, ' ===>', output_eml_file)
        anonymize_and_write_docs(input_file, output_eml_file)
        print('people:', people)

anonymize_docs()